import gradio as gr
import google.generativeai as genai
from docx import Document
from PIL import Image
import os

# 1. API Configuration using Hugging Face Secret
# In 2026, 'gemini-3-flash-preview' is the most stable high-speed model
MODEL_ID = 'gemini-3-flash-preview' 
api_key = os.getenv("GEMINI_API_KEY")

if api_key:
    genai.configure(api_key=api_key)
else:
    print("Warning: GEMINI_API_KEY not found in environment secrets.")

def process_document(input_img):
    if input_img is None:
        return None, "Error: No image uploaded.", ""
    
    if not api_key:
        return None, "Error: API Key missing in Space Secrets (GEMINI_API_KEY).", ""

    try:
        # Load the 2026 stable Flash model
        model = genai.GenerativeModel(MODEL_ID)
        
        # Convert Gradio numpy image to PIL for Gemini
        pil_img = Image.fromarray(input_img)
        
        # Expert prompt for high-fidelity document extraction
        prompt = """
        Extract all text from this document accurately. 
        - Identify titles and align them correctly.
        - Preserve Bold and Italic formatting.
        - Group lines into logical paragraphs.
        - If there are handwritten notes, transcribe them faithfully.
        """
        
        # Generate Content
        response = model.generate_content([prompt, pil_img])
        
        if not response or not response.text:
            return None, "Error: Model returned an empty response.", ""

        extracted_text = response.text

        # 2. Create Formatted Word Document
        doc = Document()
        for line in extracted_text.split('\n'):
            clean_line = line.strip()
            if clean_line:
                p = doc.add_paragraph()
                # Basic cleaning of markdown tags if Gemini adds them
                text_to_write = clean_line.replace('**', '').replace('*', '')
                run = p.add_run(text_to_write)
                if '**' in line: run.bold = True
                if '*' in line and '**' not in line: run.italic = True
        
        output_path = "Converted_Document.docx"
        doc.save(output_path)
        
        return output_path, "✅ Conversion Successful!", extracted_text

    except Exception as e:
        return None, f"❌ System Error: {str(e)}", ""

# --- Gradio UI Setup ---
with gr.Blocks(theme=gr.themes.Soft(), title="Gemini 3 Smart OCR") as demo:
    gr.Markdown("# 🖋️ AI Document Architect (Gemini 3)")
    gr.Markdown("Convert messy handwriting or document scans into formatted Word files instantly.")
    
    with gr.Row():
        with gr.Column(scale=1):
            input_image = gr.Image(label="Source Image", type="numpy")
            submit_btn = gr.Button("🚀 Convert to Word", variant="primary")
            
            # --- Added Example Images ---
            gr.Examples(
                examples=["image1.jpg", "image2.jpg"],
                inputs=input_image,
                label="Sample Notes"
            )
            
        with gr.Column(scale=1):
            status_msg = gr.Textbox(label="Status", interactive=False)
            download_link = gr.File(label="📄 Download Word File")
            
    with gr.Accordion("Review Extracted Text", open=False):
        extracted_text = gr.TextArea(label="Text Preview", lines=12)

    submit_btn.click(
        fn=process_document,
        inputs=input_image,
        outputs=[download_link, status_msg, extracted_text]
    )

if __name__ == "__main__":
    demo.launch()